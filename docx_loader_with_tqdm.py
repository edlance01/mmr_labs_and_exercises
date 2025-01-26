from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument
import os
from tqdm import tqdm


class DocxLoader():

    def process_file(self, file_path):

        # //EL TODO check figures image output folder

        doc = DocxDocument(file_path)
        text_chunks = []

        # Extract text content from paragraphs
        with tqdm(total=len(doc.paragraphs), desc="Processing Paragraphs", unit = "para") as pbar:
            for para in doc.paragraphs:
                if para.text.strip():
                    text_chunks.append(para.text)
                pbar.update(1)

        # Use RecursiveCharacterTextSplitter to chunk the text
       
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size= 500, chunk_overlap=50
        )

        with tqdm(total=len(text_chunks), desc="Splitting Text", unit="chunk") as chk:
            split_chunks=[]
            for chunk in text_chunks:
                split_chunks.extend(text_splitter.split_text(chunk))
                chk.update(1)
        return split_chunks
        
            #text_chunks_docx = text_splitter.split_text("\n".join(text_chunks))

          

        print(f"Finished:{len(split_chunks)}")

        # Wrap the method execution with a progress bar
        #file_id = MMR_DAO().store_file_and

def execute_with_progress_bar(method, steps):
    with tqdm(total=steps, desc="Processing", unit="step") as pbar:
        for _ in range(steps):
            method(1)  # Execute one step at a time
            pbar.update(1)


# Example usage
if __name__ == "__main__":
    docxloader = DocxLoader()
    docxloader.process_file("./Jupyter_Notebook_Info.docx")
